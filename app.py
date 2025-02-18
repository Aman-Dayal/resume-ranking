from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import PyPDF2
import docx
import io
from typing import List, Dict
from pydantic import BaseModel
import google.generativeai as genai
import uvicorn
import json
import pandas as pd 
from io import BytesIO
import os
from dotenv import load_dotenv
import re

load_dotenv()

app = FastAPI(
    title="Resume Ranking System",
    description="""API for extracting job requirements and ranking resumes.
    
    ## Features
    - Extract key requirements from job descriptions
    - Rank resumes against job requirements
    - Support for PDF and DOCX files
    - Integration with Gemini AI for text processing
    
    ## Error Codes
    - 400: Bad Request (invalid input)
    - 401: Unauthorized (missing/invalid API key)
    - 422: Unprocessable Entity (invalid file format)
    - 500: Internal Server Error (processing failure)
    """,
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

genai.configure(api_key=os.getenv('GENAI_API_KEY'))

model = genai.GenerativeModel('gemini-2.0-flash')

class CriteriaResponse(BaseModel):
    status_code: int
    criteria: List[str]

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file with enhanced error handling"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        if not pdf_reader.pages:
            raise ValueError("PDF contains no readable text")
        return "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
    except PyPDF2.PdfReadError:
        raise HTTPException(
            status_code=422,
            detail="Invalid PDF structure - unable to read file"
        )
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"PDF processing error: {str(e)}"
        )

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file with enhanced error handling"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        if not doc.paragraphs:
            raise ValueError("DOCX contains no readable text")
        return " ".join([paragraph.text for paragraph in doc.paragraphs])
    except docx.opc.exceptions.PackageNotFoundError:
        raise HTTPException(
            status_code=422,
            detail="Invalid DOCX file - corrupted or empty"
        )
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"DOCX processing error: {str(e)}"
        )

def extract_criteria_with_gemini(text: str) -> List[str]:
    try:
        prompt = f"""
                    Extract all key requirements from this job description such as skills, certifications, experience, and qualifications:
                    {text}
                    Format:
                        - Each requirement should be a clear, self-contained statement.
                        - No bullet points or numbers.
                        - requirement must be a quantifiable metric
                        - requirement must be a must have , no nice to haves
                    Return only the extracted requirements as a plain list.
                """
        
        response = model.generate_content(prompt)
        return response.text.strip().split('\n')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing with Gemini: {str(e)}")

def score_resume_with_gemini(resume_text: str, requirements: List[str]) -> Dict:
    """Score resume against requirements using Gemini with consistent header shortening"""
    try:        
        prompt = f"""Analyze this resume against the following job requirements. 
        For each requirement, provide Score (0-5):

        Requirements:
        {requirements}

        Resume:
        {resume_text}

        Return a valid dictionary, where:
        - The first key is 'Candidate Name' with the candidate's name.
        - Each job requirement is a key (use the exact requirement text), and its score (integer) is the value.

        Example:
        {{
            "Candidate Name": "John Doe",
            "Bachelor's degree in Computer Science": 3,
            "5+ years of Python experience": 5
        }}
        """

        response = model.generate_content(prompt)
        response_text = response.text.lstrip('```json').rstrip('```')

        try:
            result = json.loads(response_text)
            return result
   
        except json.JSONDecodeError:
            raise HTTPException(
                status_code=500,
                detail="Invalid JSON response from Gemini."
            )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Resume scoring error: {str(e)}"
        )

@app.post("/extract-criteria", 
          response_model=CriteriaResponse, 
          tags=["Extraction"],
          summary="Extract job requirements from document",
          description="""Extracts key requirements from uploaded job description documents.
          
          ## Supported Formats
          - PDF
          - DOCX
          
          ## Example Request
          ```json
          {
              "file": "job_description.pdf"
          }
          ```
          
          ## Example Response
          ```json
          {
              "criteria": [
                  "Bachelor's degree in Computer Science",
                  "5+ years of Python experience",
                  "Experience with FastAPI"
              ]
          }
          ```
          """,
          responses={
              200: {"description": "Job description processed and requirements generated sucessfully"},
              400: {"description": "Invalid file type or size"},
              422: {"description": "Unprocessable file content"},
              500: {"description": "Internal processing error"}
          })
async def extract_criteria_endpoint(file: UploadFile = File(...)) -> CriteriaResponse:
    """Extract job requirements from uploaded document"""
    if file.size > 5 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail="File size exceeds 5MB limit"
        )

    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload PDF or DOCX.")
    
    content = await file.read()
    text = extract_text_from_pdf(content) if file.content_type == "application/pdf" else extract_text_from_docx(content)
    criteria = extract_criteria_with_gemini(text)
    
    return CriteriaResponse(status_code=200,criteria=criteria)

@app.post("/rank-resumes", 
          tags=["Ranking"],
          summary="Rank resumes against job requirements",
          description="""Ranks multiple resumes against provided job requirements.
          
          ## Input
          - requirements: JSON string of job requirements
          - resumes: List of resume files (PDF/DOCX)
          
          ## Output
          - Excel file with scores for each resume
          
          ## Example Request
          ```json
          {
              "requirements": "[\"Python experience\", \"Degree in CS\"]",
              "resumes": ["resume1.pdf", "resume2.docx"]
          }
          ```
          """,
          responses={
              200: {"description": "Resumes processed, ranked and corresponding excel sheet is generated"},
              400: {"description": "Invalid input format"},
              422: {"description": "Unprocessable file content"},
              500: {"description": "Internal processing error"}
          })
async def rank_resumes(
    requirements: str = Form(...),
    resumes: List[UploadFile] = File(...)
):
    """Rank resumes against job requirements"""
    if len(resumes) > 50:
        raise HTTPException(
            status_code=400,
            detail="Maximum of 50 resumes allowed per request"
        )

    results = []
    for resume in resumes:
        if resume.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            continue
        
        content = await resume.read()
        text = extract_text_from_pdf(content) if resume.content_type == "application/pdf" else extract_text_from_docx(content)
        scores = score_resume_with_gemini(text, requirements)
        results.append(scores)
    
    df = pd.DataFrame(results)
    df['Total Score'] = df.drop('Candidate Name', axis=1).sum(axis=1)
    output = BytesIO()
    df.to_excel(output)
    output.seek(0)
    
    if not results:
        raise HTTPException(status_code=400, detail="No resumes processed successfully.")
    
    return StreamingResponse(output,status_code=200, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=resume_scores.xlsx"})

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
