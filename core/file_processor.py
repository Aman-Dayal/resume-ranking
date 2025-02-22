import io
import PyPDF2, PyPDF2.errors
import docx
from fastapi import UploadFile
import logging

# Configure logging
logging.basicConfig(filename='app.log.txt', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class BaseDoc:
    """
    Base class for handling document file processing.
    """
    
    def __init__(self, file_data: bytes):
        self.file = file_data
        self.validate_size()
        self.content = io.BytesIO(self.file)
        logging.info("Initialized BaseDoc with file size: %d bytes", len(file_data))

    @classmethod
    async def from_upload(cls, file: UploadFile):
        file_data = await file.read()
        return cls(file_data)

    def validate_size(self):
        if len(self.file) > 5 * 1024 * 1024:
            logging.error("File size exceeds 5MB limit")
            raise ValueError("File size exceeds 5MB limit")

    def extract_text(self):
        pass

class PdfDoc(BaseDoc):
    """
    Handles text extraction from PDF files.
    """

    def extract_text(self):
        logging.info("Extracting text from PDF file")
        try:
            pdf_reader = PyPDF2.PdfReader(self.content)
            if not pdf_reader.pages:
                logging.error("PDF contains no readable text")
                return {"status_code": 422, "error": "PDF contains no readable text"}

            return "\n".join([page.extract_text() or "" for page in pdf_reader.pages])

        except PyPDF2.errors.PdfReadError:
            logging.error("Invalid PDF structure - unable to read file")
            return {"status_code": 422, "error": "Invalid PDF structure - unable to read file"}

        except Exception as e:
            logging.error("PDF processing error: %s", str(e))
            return {"status_code": 422, "error": f"PDF processing error: {str(e)}"}

class DocxDoc(BaseDoc):
    """
    Handles text extraction from DOCX files.
    """

    def extract_text(self):
        logging.info("Extracting text from DOCX file")
        try:
            doc = docx.Document(self.content)
            if not doc.paragraphs:
                logging.error("DOCX contains no readable text")
                return {"status_code": 422, "error": "DOCX contains no readable text"}

            return " ".join([paragraph.text for paragraph in doc.paragraphs])

        except docx.exceptions.PythonDocxError:
            logging.error("Invalid DOCX file - corrupted or empty")
            return {"status_code": 422, "error": "Invalid DOCX file - corrupted or empty"}

        except Exception as e:
            logging.error("DOCX processing error: %s", str(e))
            return {"status_code": 422, "error": f"DOCX processing error: {str(e)}"}

class DocProcessor:
    """
    Orchestrates document processing by determining file type 
    and extracting text accordingly.
    """

    def __init__(self, file: UploadFile):
        self.file = file

    async def process(self):
        logging.info("Processing document of type: %s", self.file.content_type)
        doc = await BaseDoc.from_upload(self.file)

        if self.file.content_type == "application/pdf":
            return PdfDoc(doc.file).extract_text()

        elif self.file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return DocxDoc(doc.file).extract_text()

        else:
            logging.error("Unsupported file type: %s", self.file.content_type)
            return {"status_code": 422, "error": "Unsupported file type. Upload PDF or DOCX."}
